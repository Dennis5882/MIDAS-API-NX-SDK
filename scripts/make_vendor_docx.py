"""Render docs/vendor_report_ko.md as a .docx for handing to MIDASIT.

The markdown file stays the source of truth; this only reformats it, so the
Word document never has to be hand-edited and can be regenerated after any
correction to the report.

    pip install python-docx        # not in [dev] — this doesn't ship or run in CI
    python scripts/make_vendor_docx.py [output.docx]

Supports the markdown subset the report actually uses: ``#``/``##``/``###``
headings, paragraphs, ``- `` bullets with indented continuations, ``|`` tables,
fenced code blocks, ``>`` block quotes, ``---`` rules, and inline ``**bold**``
and ``` `code` ``` (including code nested inside bold).
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SOURCE = Path(__file__).resolve().parent.parent / "docs" / "vendor_report_ko.md"

BODY_FONT, MONO_FONT = "맑은 고딕", "Consolas"
CODE_FILL, HEAD_FILL = "F7F7F7", "EDEDED"
MUTED = RGBColor(0x66, 0x66, 0x66)
ACCENT = RGBColor(0x1F, 0x3B, 0x63)          # headings
SEVERITY = {                                  # triage colours in the 심각도 column
    "치명적": RGBColor(0xB3, 0x1B, 0x1B),
    "높음": RGBColor(0xC0, 0x5A, 0x10),
    "중간": RGBColor(0x8A, 0x6D, 0x00),
    "낮음": MUTED,
}
ZEBRA = "FAFAFA"
#: The report is short enough to navigate by its A-/B- numbering, so the
#: contents page is off. Word's navigation pane still works from the headings.
INCLUDE_TOC = False

#: Report shell. The markdown carries the content; the document identity lives
#: here, so the .md stays readable on its own and isn't cluttered with cover
#: metadata. Fill in 작성자 before sending.
TITLE = "MIDAS NX Open API 이슈 리포트"
SUBTITLE = "MIDAS CIVIL NX 2026 / GEN NX 2026 Open API 검증 결과"
DOC_INFO = [
    ("문서명", TITLE),
    ("버전", "1.1"),
    ("작성일", "2026-07-27"),
    ("수신", "MIDASIT 개발팀 / 기획팀"),
    ("검증 대상", "MIDAS CIVIL NX 2026 v2.2 (build 06/18/2026) · v2.1 (build 06/05/2026); "
                "GEN NX 2026 v2.1 (build 07/28/2026) — A-1만"),
    ("첨부", "vendor_repro_nmas.py — A-1 재현 스크립트"),
]


# --------------------------------------------------------------------------
# markdown -> block list
# --------------------------------------------------------------------------


def parse(md: str) -> list:
    """Return a list of ("kind", payload) blocks."""
    lines = md.split("\n")
    blocks, i = [], 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):                                  # fenced code
            i += 1
            buf = []
            while i < len(lines) and not lines[i].startswith("```"):
                buf.append(lines[i])
                i += 1
            blocks.append(("code", "\n".join(buf)))
            i += 1
        elif line.startswith("|"):                                  # table
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                    rows.append(cells)
                i += 1
            blocks.append(("table", rows))
        elif re.match(r"#{1,3} ", line):                            # heading
            level = len(line) - len(line.lstrip("#"))
            blocks.append((f"h{level}", line[level:].strip()))
            i += 1
        elif line.startswith(">"):                                  # block quote
            buf = []
            while i < len(lines) and lines[i].startswith(">"):
                buf.append(lines[i].lstrip(">").strip())
                i += 1
            blocks.append(("quote", " ".join(buf)))
        elif line.startswith("- ") or line.startswith("  - "):      # bullet
            indent = 1 if line.startswith("  - ") else 0
            text = line.split("- ", 1)[1]
            i += 1
            pad = "    " if indent else "  "
            while (i < len(lines) and lines[i].startswith(pad) and lines[i].strip()
                   and not lines[i].lstrip().startswith("- ")):
                text += " " + lines[i].strip()
                i += 1
            blocks.append(("bullet", (indent, text)))
        elif line.strip() == "---":
            blocks.append(("rule", ""))
            i += 1
        elif not line.strip():
            i += 1
        else:                                                       # paragraph
            buf = [line.strip()]
            i += 1
            while i < len(lines) and lines[i].strip() and not re.match(
                    r"(#{1,3} |\||```|>|- |---$)", lines[i]):
                buf.append(lines[i].strip())
                i += 1
            blocks.append(("p", " ".join(buf)))
    return blocks


# --------------------------------------------------------------------------
# low-level helpers python-docx has no API for
# --------------------------------------------------------------------------


def shade(element, fill: str) -> None:
    """Apply a solid background to a paragraph or table cell."""
    pr = element.get_or_add_tcPr() if element.tag.endswith("tc") else \
        element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pr.append(shd)


def border(paragraph, edges: dict) -> None:
    pr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge, (size, color) in edges.items():
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), color)
        pbdr.append(el)
    pr.append(pbdr)


def east_asian(run) -> None:
    """Pin the East-Asian font too, or Word falls back for the Hangul runs.

    Assigning font.name first is what materialises <w:rPr><w:rFonts>; on a run
    that has never had a font set, both are absent and the set() would fail.
    """
    name = run.font.name or BODY_FONT
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def display_width(text: str) -> int:
    """Rough rendered width in units of one Latin character.

    Sizing columns by ``len()`` is what made the first draft look wrong: a
    Korean glyph is about twice as wide as a Latin one, so a column of Korean
    prose came out under-weighted against a column of short ASCII codes.
    Markup is stripped because it isn't rendered.
    """
    plain = re.sub(r"\*\*|`", "", text)
    return sum(2 if ord(ch) > 0x2E80 else 1 for ch in plain)


def set_table_layout(table, widths, outer: str = "808080",
                     inner: str = "D9D9D9") -> None:
    """Write the column widths Word actually honours.

    ``cell.width`` alone does not survive: Word lays a fixed-layout table out
    from ``<w:tblGrid>``, which python-docx fills with equal columns at
    creation time. Both have to be set, and the layout has to be declared
    fixed or Word re-fits everything to content anyway.
    """
    tbl_pr = table._tbl.tblPr

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    borders = OxmlElement("w:tblBorders")
    for edge, (size, color) in (
        ("top", (8, outer)), ("bottom", (8, outer)),
        ("left", (2, inner)), ("right", (2, inner)),
        ("insideH", (2, inner)), ("insideV", (2, inner)),
    ):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tbl_pr.append(borders)

    margins = OxmlElement("w:tblCellMar")
    for edge, value in (("top", 60), ("bottom", 60), ("left", 108), ("right", 108)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")
        margins.append(el)
    tbl_pr.append(margins)

    for grid_col, width in zip(table._tbl.tblGrid.findall(qn("w:gridCol")), widths):
        grid_col.set(qn("w:w"), str(width.twips))


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:tblHeader"))


def tight(paragraph, before=0, after=0, exact=None) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if exact:
        pf.line_spacing = Pt(exact)


def field(paragraph, instr: str, cached: str = "") -> None:
    """A Word field (PAGE, NUMPAGES, ...) with a cached value.

    The cached text is what shows before Word recalculates, so a document that
    is opened and never refreshed still reads correctly.
    """
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instr)
    if cached:
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = cached
        run.append(text)
        fld.append(run)
    paragraph._p.append(fld)


def page_break(doc) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_cover(doc, usable) -> None:
    spacer = doc.add_paragraph()
    tight(spacer, exact=90)

    p = doc.add_paragraph()
    tight(p, after=4)
    r = p.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(26)
    east_asian(r)

    p = doc.add_paragraph()
    tight(p, after=34)
    r = p.add_run(SUBTITLE)
    r.font.size = Pt(11.5)
    r.font.color.rgb = MUTED
    east_asian(r)

    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    widths = [Cm(3.2), Cm(usable.cm - 3.2)]
    set_table_layout(table, widths, outer="BFBFBF", inner="E8E8E8")
    for label, value in DOC_INFO:
        row = table.add_row()
        for ci, text in enumerate((label, value)):
            cell = row.cells[ci]
            cell.width = widths[ci]
            cp = cell.paragraphs[0]
            tight(cp, before=2.5, after=2.5)
            add_runs(cp, text, bold=(ci == 0), size=9.5)
        shade(row.cells[0]._tc, HEAD_FILL)


def add_toc(doc, blocks) -> None:
    p = doc.add_paragraph(style="Heading 1")
    add_runs(p, "목차")
    border(p, {"bottom": (8, "999999")})
    for kind, payload in blocks:
        if kind not in ("h1", "h2"):
            continue
        if payload in (TITLE, "목차"):
            continue
        entry = doc.add_paragraph()
        tight(entry, before=1.5, after=1.5)
        entry.paragraph_format.left_indent = Cm(0.4 if kind == "h1" else 1.1)
        add_runs(entry, payload, bold=(kind == "h1"))
        if kind == "h1":
            for r in entry.runs:
                r.font.size = Pt(10.5)


def add_header_footer(doc) -> None:
    head = doc.sections[0].header.paragraphs[0]
    head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tight(head, after=0)
    r = head.add_run(TITLE)
    r.font.size = Pt(8)
    r.font.color.rgb = MUTED
    east_asian(r)
    border(head, {"bottom": (4, "D9D9D9")})

    foot = doc.sections[0].footer.paragraphs[0]
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field(foot, " PAGE ", "1")
    sep = foot.add_run(" / ")
    field(foot, " NUMPAGES ", "1")
    for r in foot.runs + [sep]:
        r.font.size = Pt(8.5)
        r.font.color.rgb = MUTED


# --------------------------------------------------------------------------
# inline runs
# --------------------------------------------------------------------------


def add_runs(paragraph, text: str, bold: bool = False, size: float | None = None):
    for token in re.split(r"(\*\*.+?\*\*|`[^`]+`)", text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            add_runs(paragraph, token[2:-2], bold=True, size=size)
        elif token.startswith("`") and token.endswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = MONO_FONT
            r.font.size = Pt((size or 9.5) - 0.7)
            r.bold = bold
            east_asian(r)
        else:
            r = paragraph.add_run(token)
            r.bold = bold
            if size:
                r.font.size = Pt(size)


# --------------------------------------------------------------------------
# build
# --------------------------------------------------------------------------


def build(blocks: list, out_path: Path) -> None:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.paragraph_format.space_after = Pt(6)
    # Hangul sits taller than Latin, so it needs more leading to stay readable.
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.widow_control = True

    for level, size, before in ((1, 16, 26), (2, 12.5, 18), (3, 10.5, 13)):
        st = doc.styles[f"Heading {level}"]
        st.font.name = BODY_FONT
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = ACCENT
        st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(7)
        # A heading stranded at the foot of a page is the most visible way a
        # generated document looks unedited.
        st.paragraph_format.keep_with_next = True
        st.paragraph_format.line_spacing = 1.25

    for section in doc.sections:
        section.page_height, section.page_width = Cm(29.7), Cm(21.0)
        section.top_margin = section.bottom_margin = Cm(2.0)
        section.left_margin = section.right_margin = Cm(2.0)
    # Length - Length yields a bare EMU int, so rebuild it as a Length.
    section = doc.sections[0]
    usable = Cm(section.page_width.cm - section.left_margin.cm
                - section.right_margin.cm)

    add_header_footer(doc)
    add_cover(doc, usable)
    page_break(doc)
    if INCLUDE_TOC:
        add_toc(doc, blocks)
        page_break(doc)

    title_seen = False
    for kind, payload in blocks:
        if kind in ("h1", "h2", "h3"):
            # The markdown's own H1 title is already on the cover.
            if kind == "h1" and not title_seen:
                title_seen = True
                continue
            p = doc.add_paragraph(style=f"Heading {kind[1]}")
            add_runs(p, payload)
            if kind == "h1":
                border(p, {"bottom": (8, "B0B8C4")})

        elif kind == "p":
            add_runs(doc.add_paragraph(), payload)

        elif kind == "bullet":
            indent, text = payload
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.6 + indent * 0.5)
            p.paragraph_format.first_line_indent = Cm(-0.35)
            tight(p, before=1, after=1)
            p.add_run("• " if not indent else "– ")
            add_runs(p, text)

        elif kind == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.55)
            p.paragraph_format.right_indent = Cm(0.25)
            tight(p, before=4, after=4)
            p.paragraph_format.line_spacing = 1.42
            shade(p._p, "F4F6F9")
            border(p, {"left": (18, "8A9BB5")})
            add_runs(p, payload)

        elif kind == "code":
            lines = payload.split("\n")
            for n, line in enumerate(lines):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.45)
                p.paragraph_format.right_indent = Cm(0.25)
                tight(p, exact=11.5)
                shade(p._p, CODE_FILL)
                edges = {}
                if n == 0:
                    edges["top"] = (4, "D9D9D9")
                if n == len(lines) - 1:
                    edges["bottom"] = (4, "D9D9D9")
                if edges:
                    border(p, edges)
                r = p.add_run(line or " ")
                r.font.name = MONO_FONT
                r.font.size = Pt(8.2)
                east_asian(r)
            tight(doc.add_paragraph(), exact=5)

        elif kind == "table":
            rows = payload
            if not rows:
                continue
            ncol = max(len(r) for r in rows)
            table = doc.add_table(rows=0, cols=ncol)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False

            # Weight by the widest cell but damp it with the column average, so
            # one long outlier doesn't starve every other column.
            widest = [max(display_width(r[c]) if c < len(r) else 0 for r in rows)
                      for c in range(ncol)]
            mean = [sum(display_width(r[c]) if c < len(r) else 0 for r in rows)
                    / len(rows) for c in range(ncol)]
            weights = [max(1.0, 0.45 * w + 0.55 * m) ** 0.72
                       for w, m in zip(widest, mean)]
            total = sum(weights)
            floor = Cm(1.5)
            widths = [max(floor, Cm(usable.cm * w / total)) for w in weights]
            overflow = sum(w.cm for w in widths) - usable.cm
            if overflow > 0:                       # give the excess back proportionally
                slack = [w for w in widths if w > floor]
                take = overflow / len(slack) if slack else 0
                widths = [Cm(max(floor.cm, w.cm - take)) if w > floor else w
                          for w in widths]
            set_table_layout(table, widths)

            for ri, cells in enumerate(rows):
                row = table.add_row()
                if ri == 0:
                    repeat_header(row)
                for ci in range(ncol):
                    cell = row.cells[ci]
                    cell.width = widths[ci]
                    p = cell.paragraphs[0]
                    tight(p, before=2.4, after=2.4)
                    p.paragraph_format.line_spacing = 1.3
                    text = cells[ci] if ci < len(cells) else ""
                    add_runs(p, text, bold=(ri == 0), size=9)
                    if ri == 0:
                        shade(cell._tc, HEAD_FILL)
                    else:
                        # Colour the triage column so severity reads at a glance.
                        colour = SEVERITY.get(text.strip())
                        if colour is not None:
                            for r in p.runs:
                                r.font.color.rgb = colour
                                r.bold = True
                        if ri % 2 == 0:
                            shade(cell._tc, ZEBRA)
            tight(doc.add_paragraph(), exact=6)

        elif kind == "rule":
            p = doc.add_paragraph()
            tight(p, before=3, after=3)
            border(p, {"bottom": (6, "BFBFBF")})

    # Footer: page numbers, so a printed copy stays in order.
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)

    doc.save(out_path)


def main() -> int:
    out_path = Path(sys.argv[1]) if len(sys.argv) > 1 else SOURCE.with_suffix(".docx")
    blocks = parse(SOURCE.read_text(encoding="utf-8"))
    build(blocks, out_path)

    kinds: dict = {}
    for k, _ in blocks:
        kinds[k] = kinds.get(k, 0) + 1
    print(f"wrote {out_path}  ({out_path.stat().st_size:,} bytes)")
    print("blocks: " + ", ".join(f"{k}={v}" for k, v in sorted(kinds.items())))
    return 0


if __name__ == "__main__":
    sys.exit(main())
